from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).parent))
import build_product_plan as base


ROOT = Path(r"D:\work\fix-pro")
SOURCE = ROOT / "docs" / "设计方案-V1.md"
OUTPUT = ROOT / "docs" / "设计方案-V1.docx"


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(base.INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Heading 1", 16, base.BLUE, 18, 10),
        ("Heading 2", 13, base.BLUE, 14, 7),
        ("Heading 3", 12, base.DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_header_footer(section) -> None:
    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("县域家居水电运维平台  |  设计方案 V1.2")
    base.set_run_font(r, size=8.5, color=base.MUTED, bold=True)

    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("产品设计基线  ·  ")
    base.set_run_font(r, size=8.5, color=base.MUTED)
    r2 = p.add_run()
    base.set_run_font(r2, size=8.5, color=base.MUTED)
    base.add_field(r2, "PAGE")


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("PRODUCT DESIGN SPEC  ·  V1.2")
    base.set_run_font(r, size=10, color=base.BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("县域家居水电运维平台")
    base.set_run_font(r, size=27, color=base.NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("服务电商小程序 · 企业维保 · 师傅工作台 · React 管理后台")
    base.set_run_font(r, size=14, color=base.DARK_BLUE)

    metadata = [
        ("上游依据", "产品方案 V1.2"),
        ("设计主题", "可信服务电商：好找、好选、可组合、可追溯"),
        ("覆盖范围", "分类搜索、购物车、多诉求、企业项目、履约端与管理后台"),
        ("首版重点", "C 端电商选购 + 约 10 万元级企业维保标杆项目"),
        ("文档状态", "研发与设计协作基线 · 2026 年 8 月"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(f"{label}：")
        base.set_run_font(r, size=10.5, color=base.NAVY, bold=True)
        r = p.add_run(value)
        base.set_run_font(r, size=10.5, color=base.INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(8)
    base.add_inline(p, "设计判断：不做低价家政撮合界面，用清楚的状态、价格边界与证据链建立县域长期信任。", size=12, color=base.NAVY)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), base.PALE_GOLD)
    p_pr.append(shd)
    doc.add_page_break()


def add_code_block(doc: Document, code_lines: list[str]) -> None:
    for idx, line in enumerate(code_lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.14)
        p.paragraph_format.right_indent = Inches(0.14)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0 if idx < len(code_lines) - 1 else 7)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line if line else " ")
        base.set_run_font(r, name="Microsoft YaHei", size=8.8, color="263238")
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F6F8FA")
        p_pr.append(shd)


def add_body(doc: Document, text: str) -> None:
    lines = text.splitlines()
    i = 0
    skipped_title = False
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped or stripped == "---":
            i += 1
            continue
        if stripped.startswith("# ") and not skipped_title:
            skipped_title = True
            i += 1
            continue
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip())
                i += 1
            i += 1
            add_code_block(doc, code_lines)
            continue
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.right_indent = Inches(0.18)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(10)
            base.add_inline(p, stripped[2:], size=11, color=base.NAVY)
            p_pr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), base.PALE_BLUE)
            p_pr.append(shd)
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            rows, i = base.parse_table(lines, i)
            base.add_markdown_table(doc, rows)
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            base.add_inline(p, stripped[4:], size=13, color=base.BLUE)
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            base.add_inline(p, stripped[3:], size=16, color=base.BLUE)
            i += 1
            continue
        if re.match(r"^- ", stripped):
            p = doc.add_paragraph(style="List Bullet")
            base.add_inline(p, stripped[2:], size=11)
            i += 1
            continue
        if re.match(r"^\d+\. ", stripped):
            p = doc.add_paragraph(style="List Number")
            base.add_inline(p, re.sub(r"^\d+\. ", "", stripped), size=11)
            i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt == "---" or nxt.startswith(("#", ">", "|", "- ", "```")) or re.match(r"^\d+\. ", nxt):
                break
            paragraph_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        base.add_inline(p, " ".join(paragraph_lines), size=11)


def build() -> None:
    doc = Document()
    configure(doc)
    add_header_footer(doc.sections[0])
    add_cover(doc)
    add_body(doc, SOURCE.read_text(encoding="utf-8"))
    props = doc.core_properties
    props.title = "县域家居水电运维平台｜设计方案 V1.2"
    props.subject = "客户微信小程序、师傅工作台与 React 管理后台 UI/UX 设计方案"
    props.author = "创业项目产品与研发团队"
    props.keywords = "UI, UX, 微信小程序, React, 管理后台, 工单, 报价, 证据, 设计系统"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
