from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).parent))
import build_product_plan as base


ROOT = Path(r"D:\work\fix-pro")
SOURCE = ROOT / "docs" / "技术方案-V1.md"
OUTPUT = ROOT / "docs" / "技术方案-V1.docx"


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(base.INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Heading 1", 16, base.BLUE, 18, 10),
        ("Heading 2", 13, base.BLUE, 14, 7),
        ("Heading 3", 12, base.DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_header_footer(section) -> None:
    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("县域家居水电运维平台  |  技术方案 V1.2")
    base.set_run_font(r, size=8.5, color=base.MUTED, bold=True)

    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("开发基线讨论稿  ·  ")
    base.set_run_font(r, size=8.5, color=base.MUTED)
    r2 = p.add_run()
    base.set_run_font(r2, size=8.5, color=base.MUTED)
    base.add_field(r2, "PAGE")


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("TECHNICAL DESIGN  ·  V1.2")
    base.set_run_font(r, size=10, color=base.BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("县域家居水电运维平台")
    base.set_run_font(r, size=27, color=base.NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("生产级软件技术方案")
    base.set_run_font(r, size=15, color=base.DARK_BLUE)

    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "技术维度"
    table.cell(0, 1).text = "决策基线"
    for cell in table.rows[0].cells:
        base.set_cell_shading(cell, base.NAVY)
        for run in cell.paragraphs[0].runs:
            base.set_run_font(run, size=9.5, color=base.WHITE, bold=True)
    base.set_repeat_table_header(table.rows[0])

    rows = [
        ("后端", "Java LTS + Spring Boot 3 + Spring Security + MyBatis"),
        ("管理后台", "React + TypeScript + Vite + Ant Design"),
        ("微信小程序", "微信原生小程序 + TypeScript"),
        ("数据", "MySQL + Redis + 私有对象存储 + Flyway"),
        ("架构", "模块化单体 + Outbox + 独立后台任务，渐进拆分"),
    ]
    for i, (label, value) in enumerate(rows, 1):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        base.set_cell_shading(table.cell(i, 0), base.PALE_BLUE)
        for run in table.cell(i, 0).paragraphs[0].runs:
            base.set_run_font(run, size=9.5, color=base.NAVY, bold=True)
        for run in table.cell(i, 1).paragraphs[0].runs:
            base.set_run_font(run, size=9.5, color=base.INK)
    base.set_table_geometry(table, [1900, 7460])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(8)
    base.add_inline(
        p,
        "首版同时守住服务电商结算与拆单正确性、企业项目履约、证据可信度和权限正确性。",
        size=12.5,
        color=base.NAVY,
    )
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), base.PALE_GOLD)
    p_pr.append(shd)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("上游依据：产品方案 V1.2  |  2026 年 8 月")
    base.set_run_font(r, size=9, color=base.MUTED)
    doc.add_page_break()


def add_code_block(doc: Document, code_lines: list[str], language: str) -> None:
    if language:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(language.upper())
        base.set_run_font(r, name="Consolas", size=8, color=base.BLUE, bold=True)
    for idx, line in enumerate(code_lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.right_indent = Inches(0.12)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0 if idx < len(code_lines) - 1 else 6)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line if line else " ")
        base.set_run_font(r, name="Consolas", size=8.2, color="263238")
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F6F8FA")
        p_pr.append(shd)


def add_body(doc: Document, text: str) -> None:
    lines = text.splitlines()
    i = 0
    skipped_title = False
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped or stripped == "---":
            i += 1
            continue
        if stripped.startswith("# ") and not skipped_title:
            skipped_title = True
            i += 1
            continue
        if stripped.startswith("```"):
            language = stripped[3:].strip()
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip())
                i += 1
            i += 1
            add_code_block(doc, code_lines, language)
            continue
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.right_indent = Inches(0.18)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(10)
            base.add_inline(p, stripped[2:], size=11, color=base.NAVY)
            p_pr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), base.PALE_BLUE)
            p_pr.append(shd)
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            rows, i = base.parse_table(lines, i)
            base.add_markdown_table(doc, rows)
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            base.add_inline(p, stripped[4:], size=12, color=base.DARK_BLUE)
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            base.add_inline(p, stripped[3:], size=13, color=base.BLUE)
            i += 1
            continue
        if stripped.startswith("# "):
            p = doc.add_paragraph(style="Heading 1")
            base.add_inline(p, stripped[2:], size=16, color=base.BLUE)
            i += 1
            continue
        if re.match(r"^- ", stripped):
            p = doc.add_paragraph(style="List Bullet")
            base.add_inline(p, stripped[2:], size=11)
            i += 1
            continue
        if re.match(r"^\d+\. ", stripped):
            p = doc.add_paragraph(style="List Number")
            base.add_inline(p, re.sub(r"^\d+\. ", "", stripped), size=11)
            i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (
                not nxt
                or nxt == "---"
                or nxt.startswith(("#", ">", "|", "- ", "```"))
                or re.match(r"^\d+\. ", nxt)
            ):
                break
            paragraph_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        base.add_inline(p, " ".join(paragraph_lines), size=11)


def build() -> None:
    doc = Document()
    configure(doc)
    add_header_footer(doc.sections[0])
    add_cover(doc)
    add_body(doc, SOURCE.read_text(encoding="utf-8"))
    props = doc.core_properties
    props.title = "县域家居水电运维平台｜技术方案 V1.2"
    props.subject = "Java Spring Boot、React 管理后台与微信小程序生产级技术设计"
    props.author = "创业项目研发团队"
    props.keywords = "Java, Spring Boot, React, 微信小程序, 工单, 支付, 影像证据, MySQL"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
