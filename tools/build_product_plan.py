from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\work\fix-pro")
SOURCE = ROOT / "docs" / "产品方案-V1.md"
OUTPUT = ROOT / "docs" / "产品方案-V1.docx"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "202124"
MUTED = "667085"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF8E8"
BORDER = "D0D5DD"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="Microsoft YaHei", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_all_runs(paragraph, size=10.5, color=INK):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def add_inline(paragraph, text: str, size=10.5, color=INK) -> None:
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=max(size - 0.5, 9), color=DARK_BLUE)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, color=color)


def add_field(run, instruction: str) -> None:
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def set_paragraph_keep(paragraph, keep_next=False, keep_lines=True):
    paragraph.paragraph_format.keep_with_next = keep_next
    paragraph.paragraph_format.keep_together = keep_lines


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.38)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 16, 7),
        ("Heading 2", 13, BLUE, 12, 5),
        ("Heading 3", 11.5, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.page_break_before = False

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.34)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.18


def add_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("县域家居水电运维平台  |  产品方案 V1.2")
    set_run_font(r, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("创始人讨论稿  ·  ")
    set_run_font(r, size=8.5, color=MUTED)
    r2 = p.add_run()
    set_run_font(r2, size=8.5, color=MUTED)
    add_field(r2, "PAGE")


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(50)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("PRODUCT BRIEF  ·  V1.2")
    set_run_font(r, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("县域家居水电运维平台")
    set_run_font(r, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("服务电商小程序 + 企业维保 + 师傅工作台 + Web 中后台")
    set_run_font(r, size=14, color=DARK_BLUE)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "方案维度"
    table.cell(0, 1).text = "方案结论"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=9.5, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("产品定位", "以可信履约工单为核心的县域家居运维数字化经营系统"),
        ("首发客户", "散户业主；物业与政企作为合同工单能力补充"),
        ("首版目标", "可搜索、可分类、可购物车、多诉求可拆单、企业项目可闭环"),
        ("建议策略", "先单城自营验证，再扩展 B 端与加盟能力"),
    ]
    for i, (label, value) in enumerate(rows, 1):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        set_cell_shading(table.cell(i, 0), PALE_BLUE)
        for run in table.cell(i, 0).paragraphs[0].runs:
            set_run_font(run, size=9.5, color=NAVY, bold=True)
        for run in table.cell(i, 1).paragraphs[0].runs:
            set_run_font(run, size=9.5, color=INK)
    set_table_geometry(table, [1800, 7560])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("产品主张")
    set_run_font(r, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.right_indent = Inches(0.22)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    add_inline(p, "让客户敢下单、让公司管得住、让师傅干得清、让每次施工有证据、让一次上门持续产生后续价值。", size=13, color=NAVY)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PALE_GOLD)
    p_pr.append(shd)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("依据：合伙人业务输入全量材料  |  2026 年 8 月")
    set_run_font(r, size=9, color=MUTED)
    doc.add_page_break()


def table_widths(headers: list[str]) -> list[int]:
    n = len(headers)
    total = 9360
    if n == 2:
        return [2100, 7260]
    if n == 3:
        return [1700, 3580, 4080]
    if n == 4:
        return [1500, 2700, 2500, 2660]
    return [total // n] * (n - 1) + [total - (total // n) * (n - 1)]


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    headers = rows[0]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for j, text in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ""
        add_inline(cell.paragraphs[0], text, size=9, color=WHITE)
        for run in cell.paragraphs[0].runs:
            run.bold = True
        set_cell_shading(cell, NAVY)
    set_repeat_table_header(table.rows[0])
    for i, row in enumerate(rows[1:], 1):
        cells = table.add_row().cells
        for j, text in enumerate(row):
            cells[j].text = ""
            add_inline(cells[j].paragraphs[0], text, size=8.8, color=INK)
            if i % 2 == 0:
                set_cell_shading(cells[j], LIGHT)
    set_table_geometry(table, table_widths(headers))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)


def parse_table(lines: list[str], start: int):
    raw = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw.append(lines[i].strip())
        i += 1
    rows = []
    for idx, line in enumerate(raw):
        parts = [p.strip() for p in line.strip("|").split("|")]
        if idx == 1 and all(re.fullmatch(r":?-{3,}:?", p) for p in parts):
            continue
        rows.append(parts)
    return rows, i


def add_body_from_markdown(doc: Document, text: str) -> None:
    lines = text.splitlines()
    i = 0
    skipped_title = False
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped or stripped == "---":
            i += 1
            continue
        if stripped.startswith("# ") and not skipped_title:
            skipped_title = True
            i += 1
            continue
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.right_indent = Inches(0.18)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(10)
            add_inline(p, stripped[2:], size=11, color=NAVY)
            p_pr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), PALE_BLUE)
            p_pr.append(shd)
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            rows, i = parse_table(lines, i)
            add_markdown_table(doc, rows)
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, stripped[4:], size=11.5, color=DARK_BLUE)
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, stripped[3:], size=13, color=BLUE)
            i += 1
            continue
        if stripped.startswith("# "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, stripped[2:], size=16, color=NAVY)
            i += 1
            continue
        if re.match(r"^- ", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
            set_paragraph_keep(p, keep_lines=False)
            i += 1
            continue
        if re.match(r"^\d+\. ", stripped):
            p = doc.add_paragraph(style="List Number")
            content = re.sub(r"^\d+\. ", "", stripped)
            add_inline(p, content)
            set_paragraph_keep(p, keep_lines=False)
            i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt == "---" or nxt.startswith(("#", ">", "|", "- ")) or re.match(r"^\d+\. ", nxt):
                break
            paragraph_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(paragraph_lines))
        set_paragraph_keep(p, keep_lines=False)


def build() -> None:
    doc = Document()
    configure_document(doc)
    add_header_footer(doc.sections[0])
    add_cover(doc)
    add_body_from_markdown(doc, SOURCE.read_text(encoding="utf-8"))

    props = doc.core_properties
    props.title = "县域家居水电运维平台｜产品方案 V1.2"
    props.subject = "微信小程序、师傅工作台与 Web 中后台产品方案"
    props.author = "创业团队"
    props.keywords = "家居运维, 微信小程序, 工单, 物业, 政企, 材料溯源"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
